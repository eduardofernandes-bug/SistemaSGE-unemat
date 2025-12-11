# -*- coding: utf-8 -*-
"""
documento.py - Gerador de documentos (DOCX e XLSX) do SGE
Gera: Plano de Atividades e Ficha de Atividades
"""

import os
import logging
import openpyxl
import openpyxl.utils
from datetime import datetime
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from openpyxl.cell.cell import MergedCell

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


class GeradorDocumentos:
    """
    Classe para gerar documentos automatizados do sistema de estágios.
    
    Conceitos POO aplicados:
    - Encapsulamento: Métodos privados para processamento interno
    - Abstração: Interface simples para gerar documentos
    - Métodos de classe: Factory methods para diferentes tipos de documentos
    """
    
    # Caminhos dos templates
    TEMPLATES_DIR = 'documentos_templates'
    OUTPUT_DIR = 'documentos_gerados'
    
    TEMPLATE_PLANO = os.path.join(TEMPLATES_DIR, 'anexo_VI_Plano_de_Atividades.docx')
    TEMPLATE_FICHA = os.path.join(TEMPLATES_DIR, 'Anexo_VII_Ficha_Atividades.xlsx')
    
    def __init__(self):
        """Inicializa o gerador e cria pasta de saída se não existir"""
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)
        os.makedirs(self.TEMPLATES_DIR, exist_ok=True)
    
    # ==================== GERAÇÃO DE PLANO DE ATIVIDADES (DOCX) ====================
    
    @classmethod
    def gerar_plano_atividades(cls, dados_estagiario, dados_empresa, cronograma, cidade=""):
        """
        Gera o Anexo VI - Plano de Atividades em DOCX.
        
        Args:
            dados_estagiario (dict): {'nome': str, 'telefone': str, 'email': str}
            dados_empresa (dict): {'nome': str, 'telefone': str, 'email': str}
            cronograma (list): [{'periodo': str, 'atividades': str}, ...]
            cidade (str): Nome da cidade
            
        Returns:
            str: Caminho do arquivo gerado
        """
        try:
            # Carrega template
            doc = Document(cls.TEMPLATE_PLANO)
            
            # Substitui placeholders nas tabelas
            for table in doc.tables:
                # Tabela de identificação (primeira tabela)
                if len(table.rows) >= 3:
                    # Estagiário
                    cls._substituir_texto_celula(table.rows[1].cells[0], 
                                                  'Digite seu nome aqui', 
                                                  dados_estagiario['nome'])
                    cls._substituir_texto_celula(table.rows[2].cells[0], 
                                                  'Digite o número de contato aqui', 
                                                  dados_estagiario['telefone'])
                    cls._substituir_texto_celula(table.rows[3].cells[0], 
                                                  'Digite seu e-mail aqui', 
                                                  dados_estagiario['email'])
                    
                    # Empresa
                    cls._substituir_texto_celula(table.rows[1].cells[1], 
                                                  'Digite o nome aqui', 
                                                  dados_empresa['nome'])
                    cls._substituir_texto_celula(table.rows[2].cells[1], 
                                                  'Digite o número de contato aqui', 
                                                  dados_empresa['telefone'])
                    cls._substituir_texto_celula(table.rows[3].cells[1], 
                                                  'Digite o e-mail aqui', 
                                                  dados_empresa['email'])
                
                # Tabela de cronograma
                if 'PERÍODO' in table.rows[0].cells[0].text:
                    # Remove linhas template (mantém cabeçalho)
                    for _ in range(len(table.rows) - 1):
                        if len(table.rows) > 1:
                            table._element.remove(table.rows[-1]._element)
                    
                    # Adiciona cronograma
                    for item in cronograma:
                        row = table.add_row()
                        row.cells[0].text = item['periodo']
                        row.cells[1].text = item['atividades']
                
                # Tabela de assinaturas
                if 'Acadêmico(a)/Estagiário(a)' in table.rows[-1].cells[0].text:
                    cls._substituir_texto_celula(table.rows[-2].cells[0], 
                                                  'Digite o nome aqui', 
                                                  dados_estagiario['nome'])
                    cls._substituir_texto_celula(table.rows[-2].cells[1], 
                                                  'Digite o nome aqui', 
                                                  dados_empresa.get('supervisor', ''))
            
            # Substitui data/cidade
            data_atual = datetime.now().strftime('%d/%m/%Y')
            for paragraph in doc.paragraphs:
                if 'Digite a cidade' in paragraph.text:
                    paragraph.text = f"{cidade}, {data_atual}"
            
            # Salva arquivo
            nome_arquivo = f"Plano_Atividades_{dados_estagiario['nome'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            caminho_saida = os.path.join(cls.OUTPUT_DIR, nome_arquivo)
            doc.save(caminho_saida)
            
            logger.info(f"Plano de Atividades gerado: {caminho_saida}")
            return caminho_saida
            
        except Exception as e:
            logger.exception(f"Erro ao gerar Plano de Atividades: {e}")
            raise
    
    @staticmethod
    def _substituir_texto_celula(celula, texto_antigo, texto_novo):
        """Substitui texto em uma célula da tabela"""
        for paragraph in celula.paragraphs:
            if texto_antigo in paragraph.text:
                paragraph.text = paragraph.text.replace(texto_antigo, texto_novo)
    
    # ==================== GERAÇÃO DE FICHA DE ATIVIDADES (XLSX) ====================
    
    @classmethod
    def gerar_ficha_atividades(cls, dados_estagiario, dados_empresa, mes, ano, atividades_diarias):
        """
        Gera o Anexo VII - Ficha de Atividades em XLSX.
        """
        try:
            wb = load_workbook(cls.TEMPLATE_FICHA)
            ws = wb.active
            
            # === PRESERVA IMAGENS (ADICIONE ISTO) ===
            imagens_originais = []
            if hasattr(ws, '_images'):
                for img in ws._images:
                    imagens_originais.append(img)
            
            # === PREENCHE CABEÇALHO (Acadêmico, Empresa, Mês, Ano) ===
            
            # Acadêmico: preenche D13
            ws['D13'] = dados_estagiario['nome']
            
            # Empresa/Instituição: preenche D14
            ws['D14'] = dados_empresa['nome']
            
            # Mês: preenche C17
            ws['C17'] = mes
            
            # Ano: preenche E17
            ws['E17'] = ano
            
            # === PREENCHE ATIVIDADES DIÁRIAS ===
            # Tabela começa na linha 20 (cabeçalho em linha 19)
            # Colunas: A=Dia, B=Início, C=Término, D=Descrição, E=Total Horas
            
            for atividade in atividades_diarias:
                dia = atividade['dia']
                
                # Válida o dia (1-31)
                if 1 <= dia <= 31:
                    linha = 19 + dia  # Linha 20 é dia 1, linha 21 é dia 2, etc
                    
                    # Coluna B: Horário Início
                    if atividade.get('inicio'):
                        ws[f'B{linha}'] = atividade['inicio']
                    
                    # Coluna C: Horário Término
                    if atividade.get('termino'):
                        ws[f'C{linha}'] = atividade['termino']
                    
                    # Coluna D: Descrição das atividades (IMPORTANTE!)
                    if atividade.get('descricao'):
                        ws[f'D{linha}'] = atividade['descricao']
                    
                    # Coluna E: Total Horas (já tem fórmula, mas pode atualizar se necessário)
                    # A fórmula já existe: =(C{linha}-B{linha})*24
            
            # === RESTAURA IMAGENS (ADICIONE ISTO) ===
            ws._images = imagens_originais
            
            # === SALVA ARQUIVO ===
            nome_arquivo = f"Ficha_Atividades_{dados_estagiario['nome'].replace(' ', '_')}_{mes}_{ano}.xlsx"
            caminho_saida = os.path.join(cls.OUTPUT_DIR, nome_arquivo)
            wb.save(caminho_saida)
            
            logger.info(f"Ficha de Atividades gerada com sucesso: {caminho_saida}")
            return caminho_saida
            
        except Exception as e:
            logger.exception(f"Erro ao gerar Ficha de Atividades: {e}")
            raise

    
    @staticmethod
    def _calcular_horas(hora_inicio, hora_termino):
        """
        Calcula diferença de horas entre dois horários.
        
        Args:
            hora_inicio (str): Formato 'HH:MM' ou 'HH:MM:SS'
            hora_termino (str): Formato 'HH:MM' ou 'HH:MM:SS'
            
        Returns:
            float: Diferença em horas
        """
        try:
            # Converte strings para datetime
            fmt = '%H:%M:%S' if ':' in hora_inicio and hora_inicio.count(':') == 2 else '%H:%M'
            
            inicio = datetime.strptime(hora_inicio, fmt)
            termino = datetime.strptime(hora_termino, fmt)
            
            # Calcula diferença
            diferenca = termino - inicio
            horas = diferenca.total_seconds() / 3600
            
            return round(horas, 2)
            
        except Exception as e:
            logger.warning(f"Erro ao calcular horas: {e}")
            return 0
    
    # ==================== MÉTODOS AUXILIARES ====================
    
    @classmethod
    def verificar_templates(cls):
        """
        Verifica se os templates existem.
        
        Returns:
            dict: {'plano': bool, 'ficha': bool}
        """
        return {
            'plano': os.path.exists(cls.TEMPLATE_PLANO),
            'ficha': os.path.exists(cls.TEMPLATE_FICHA)
        }
    
    @classmethod
    def listar_documentos_gerados(cls):
        """
        Lista todos os documentos gerados.
        
        Returns:
            list: Lista de dicionários com informações dos arquivos
        """
        if not os.path.exists(cls.OUTPUT_DIR):
            return []
        
        documentos = []
        for arquivo in os.listdir(cls.OUTPUT_DIR):
            caminho_completo = os.path.join(cls.OUTPUT_DIR, arquivo)
            if os.path.isfile(caminho_completo):
                info = {
                    'nome': arquivo,
                    'caminho': caminho_completo,
                    'tamanho': os.path.getsize(caminho_completo),
                    'data_criacao': datetime.fromtimestamp(os.path.getctime(caminho_completo))
                }
                documentos.append(info)
        
        return sorted(documentos, key=lambda x: x['data_criacao'], reverse=True)
